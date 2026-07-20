"""
文档加载与切分 - 支持PDF/DOCX/TXT/CSV
"""
import os
from typing import List, Tuple, Optional
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.config import settings


def load_pdf(file_path: str) -> str:
    """加载PDF文件"""
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n\n".join(text_parts)


def load_docx(file_path: str) -> str:
    """加载Word文档"""
    from docx import Document
    doc = Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    # 也读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))
    return "\n".join(text_parts)


def load_txt(file_path: str) -> str:
    """加载文本文件"""
    encodings = ["utf-8", "gbk", "gb2312", "utf-16"]
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别文件编码: {file_path}")


def load_csv(file_path: str) -> str:
    """加载CSV文件，转为文本"""
    import pandas as pd
    df = pd.read_csv(file_path, encoding="utf-8")
    text_parts = []
    # 表头信息
    text_parts.append(f"数据表包含以下字段: {', '.join(df.columns.tolist())}")
    text_parts.append(f"共 {len(df)} 条记录\n")
    # 每行数据
    for idx, row in df.iterrows():
        row_text = "，".join([f"{col}: {val}" for col, val in row.items()])
        text_parts.append(f"第{idx + 1}条: {row_text}")
    return "\n".join(text_parts)


def load_document(file_path: str, file_type: str) -> str:
    """
    根据文件类型加载文档内容
    """
    loaders = {
        "pdf": load_pdf,
        "docx": load_docx,
        "txt": load_txt,
        "csv": load_csv,
    }

    loader = loaders.get(file_type.lower())
    if not loader:
        raise ValueError(f"不支持的文件类型: {file_type}，支持: {list(loaders.keys())}")

    return loader(file_path)


def split_text(text: str, chunk_size: int = None, chunk_overlap: int = None) -> List[str]:
    """
    递归字符切分
    按段落→句子→字符逐级切分，保持语义完整性
    """
    chunk_size = chunk_size or settings.RAG_CHUNK_SIZE
    chunk_overlap = chunk_overlap or settings.RAG_CHUNK_OVERLAP

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        length_function=len,
    )
    return splitter.split_text(text)


def process_file(file_path: str, file_type: str) -> Tuple[str, List[str]]:
    """
    处理文件：加载 + 切分
    返回: (原始文本, 切分后的片段列表)
    """
    # 加载
    raw_text = load_document(file_path, file_type)

    # 清洗
    cleaned_text = clean_text(raw_text)

    # 切分
    chunks = split_text(cleaned_text)

    return cleaned_text, chunks


def clean_text(text: str) -> str:
    """文本清洗"""
    import re
    # 去除多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 去除多余空格
    text = re.sub(r' {2,}', ' ', text)
    # 去除首尾空白
    text = text.strip()
    return text
